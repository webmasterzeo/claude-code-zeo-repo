#!/usr/bin/env python3
"""Convert Markdown content to styled DOCX for Google Drive."""

import argparse
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def parse_markdown(md_text):
    """Parse markdown into structured blocks."""
    blocks = []
    lines = md_text.strip().split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith('### '):
            blocks.append(('h3', line[4:].strip()))
        elif line.startswith('## '):
            blocks.append(('h2', line[3:].strip()))
        elif line.startswith('# '):
            blocks.append(('h1', line[2:].strip()))
        elif line.startswith('- ') or line.startswith('* '):
            blocks.append(('list_item', line[2:].strip()))
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            blocks.append(('list_item_ordered', text))
        elif line.strip() == '':
            pass
        else:
            para_lines = [line]
            while (i + 1 < len(lines)
                   and lines[i + 1].strip()
                   and not lines[i + 1].startswith('#')
                   and not lines[i + 1].startswith('- ')
                   and not lines[i + 1].startswith('* ')
                   and not re.match(r'^\d+\.\s', lines[i + 1])):
                i += 1
                para_lines.append(lines[i])
            blocks.append(('paragraph', ' '.join(para_lines)))
        i += 1
    return blocks


def add_hyperlink(paragraph, url, text):
    """Add a real clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline_formatting(paragraph, text):
    """Add runs with real hyperlinks parsed from markdown [text](url) syntax."""
    pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    last_end = 0

    for match in re.finditer(pattern, text):
        before = text[last_end:match.start()]
        if before:
            paragraph.add_run(before)

        link_text = match.group(1)
        link_url = match.group(2)
        add_hyperlink(paragraph, link_url, link_text)

        last_end = match.end()

    remaining = text[last_end:]
    if remaining:
        paragraph.add_run(remaining)


def add_keyword_table(doc, keywords_json):
    """Add a keyword tracking table to the document."""
    keywords = json.loads(keywords_json)
    if not keywords:
        return

    para = doc.add_paragraph()
    run = para.add_run('Keyword tracking')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.bold = True

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    headers = ['Keyword', 'Type', 'Volume', 'Count']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.bold = True

    for kw in keywords:
        row = table.add_row()
        row.cells[0].text = kw.get('keyword', '')
        row.cells[1].text = kw.get('type', '')
        row.cells[2].text = str(kw.get('volume', ''))
        row.cells[3].text = str(kw.get('count', ''))
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()


def create_docx(blocks, meta_title, meta_description, output_path,
                keywords_json=None):
    """Create a styled DOCX from parsed blocks."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    if meta_title or meta_description:
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if meta_title:
            run = meta_para.add_run(f'Meta Title: {meta_title}\n')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        if meta_description:
            run = meta_para.add_run(f'Meta Description: {meta_description}')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    if keywords_json:
        add_keyword_table(doc, keywords_json)
    elif meta_title or meta_description:
        doc.add_paragraph()

    for block_type, content in blocks:
        if block_type == 'h1':
            doc.add_heading(content, level=1)
        elif block_type == 'h2':
            doc.add_heading(content, level=2)
        elif block_type == 'h3':
            doc.add_heading(content, level=3)
        elif block_type == 'paragraph':
            para = doc.add_paragraph()
            add_inline_formatting(para, content)
        elif block_type in ('list_item', 'list_item_ordered'):
            style_name = 'List Bullet' if block_type == 'list_item' else 'List Number'
            para = doc.add_paragraph(style=style_name)
            add_inline_formatting(para, content)

    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser(description='Convert Markdown to DOCX')
    parser.add_argument('input', help='Input markdown file')
    parser.add_argument('output', help='Output DOCX file')
    parser.add_argument('--meta-title', default='', help='Meta title')
    parser.add_argument('--meta-description', default='', help='Meta description')
    parser.add_argument('--keywords', default='', help='Keyword tracking JSON array')
    args = parser.parse_args()

    md_text = Path(args.input).read_text(encoding='utf-8')
    blocks = parse_markdown(md_text)
    create_docx(blocks, args.meta_title, args.meta_description, args.output,
                keywords_json=args.keywords or None)
    print(f'Created: {args.output}')


if __name__ == '__main__':
    main()
